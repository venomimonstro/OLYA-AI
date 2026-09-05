from __future__ import annotations

import hashlib
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt
from pypdf import PdfReader
from PIL import Image, ImageChops


class DocumentBuildError(RuntimeError):
    pass


class DocumentQAError(RuntimeError):
    pass


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def safe_doc_name(name: str) -> str:
    value = Path(name.replace("\\", "/")).name.strip() or "document.docx"
    if not value.lower().endswith(".docx"):
        value += ".docx"
    return value[:240]


def build_docx(spec: dict[str, Any], destination: Path) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    for level in range(1, 5):
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"

    title = str(spec.get("title") or "").strip()
    if title:
        document.add_heading(title, level=0)

    block_count = 0
    table_count = 0
    for raw in spec.get("blocks") or []:
        kind = raw.get("type")
        if kind == "heading":
            document.add_heading(str(raw.get("text") or ""), level=max(1, min(4, int(raw.get("level") or 1))))
        elif kind == "paragraph":
            document.add_paragraph(str(raw.get("text") or ""))
        elif kind in {"bullet_list", "numbered_list"}:
            style = "List Bullet" if kind == "bullet_list" else "List Number"
            for item in raw.get("items") or []:
                document.add_paragraph(str(item), style=style)
        elif kind == "table":
            rows = raw.get("rows") or []
            if not rows:
                continue
            width = max(len(row) for row in rows)
            if width < 1:
                continue
            table = document.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(width):
                    table.cell(r_idx, c_idx).text = str(row[c_idx]) if c_idx < len(row) else ""
            table_count += 1
        elif kind == "page_break":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            raise DocumentBuildError(f"Unsupported document block: {kind}")
        block_count += 1

    document.save(destination)
    if not destination.is_file() or destination.stat().st_size < 1000:
        raise DocumentBuildError("DOCX generation produced an invalid or empty file")
    return {"block_count": block_count, "table_count": table_count, "size_bytes": destination.stat().st_size}


def structural_qa(docx_path: Path, spec: dict[str, Any]) -> dict[str, Any]:
    issues: list[dict[str, Any]] = []
    try:
        doc = Document(docx_path)
    except Exception as exc:
        raise DocumentQAError("DOCX cannot be opened") from exc

    expected_tables = sum(1 for block in spec.get("blocks") or [] if block.get("type") == "table" and block.get("rows"))
    if len(doc.tables) != expected_tables:
        issues.append({"code": "table_count_mismatch", "expected": expected_tables, "actual": len(doc.tables)})

    non_empty_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    if not non_empty_text and not doc.tables:
        issues.append({"code": "empty_document"})

    unresolved = [token for token in ("TODO", "TBD", "FIXME", "{{", "}}") if token in non_empty_text]
    if unresolved:
        issues.append({"code": "unresolved_placeholders", "tokens": unresolved})

    return {"status": "passed" if not issues else "failed", "issues": issues, "paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}


def _office_binary() -> str:
    binary = shutil.which("libreoffice") or shutil.which("soffice")
    if not binary:
        raise DocumentQAError("LibreOffice is unavailable for document rendering")
    return binary


def render_docx_to_pdf(docx_path: Path, output_dir: Path, *, timeout_seconds: int = 60) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="x1-lo-") as profile:
        cmd = [_office_binary(), "--headless", "--nologo", "--nodefault", "--nofirststartwizard", f"-env:UserInstallation=file://{profile}", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)]
        env = {**os.environ, "HOME": profile}
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout_seconds, env=env, check=False)
        except subprocess.TimeoutExpired as exc:
            raise DocumentQAError("Document rendering timed out") from exc
    pdf = output_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not pdf.is_file() or pdf.stat().st_size < 500:
        message = (result.stderr or result.stdout or "LibreOffice conversion failed")[-1000:]
        raise DocumentQAError(message)
    return pdf


def render_qa(pdf_path: Path, raster_dir: Path | None = None, *, max_pages: int = 300) -> dict[str, Any]:
    try:
        reader = PdfReader(str(pdf_path))
    except Exception as exc:
        raise DocumentQAError("Rendered PDF cannot be opened") from exc
    page_count = len(reader.pages)
    if page_count < 1:
        raise DocumentQAError("Rendered PDF contains no pages")
    if page_count > max_pages:
        raise DocumentQAError("Rendered PDF exceeds QA page limit")

    issues: list[dict[str, Any]] = []
    pages: list[dict[str, Any]] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        page_info = {"page": index, "text_chars": len(text), "width_pt": round(width, 2), "height_pt": round(height, 2)}
        if width <= 0 or height <= 0:
            issues.append({"code": "invalid_page_geometry", "page": index})
        resources = page.get("/Resources")
        has_xobject = bool(resources and resources.get("/XObject"))
        page_info["has_xobject"] = has_xobject
        if not text and not has_xobject:
            issues.append({"code": "blank_page", "page": index})
        pages.append(page_info)

    raster_status = "not_checked"
    if raster_dir is not None and shutil.which("pdftoppm"):
        raster_dir.mkdir(parents=True, exist_ok=True)
        prefix = raster_dir / "page"
        cmd = ["pdftoppm", "-png", "-r", "72", str(pdf_path), str(prefix)]
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=90, check=False)
        except subprocess.TimeoutExpired as exc:
            raise DocumentQAError("Page rasterization timed out") from exc
        pngs = sorted(raster_dir.glob("page-*.png"))
        if result.returncode != 0 or len(pngs) != page_count:
            issues.append({"code": "page_rasterization_failed", "expected": page_count, "actual": len(pngs)})
            raster_status = "failed"
        else:
            raster_status = "passed"
            for idx, image in enumerate(pngs):
                pages[idx]["raster_sha256"] = sha256_file(image)
                pages[idx]["raster_bytes"] = image.stat().st_size
                with Image.open(image).convert("RGB") as page_image:
                    white = Image.new("RGB", page_image.size, "white")
                    diff = ImageChops.difference(page_image, white)
                    bbox = diff.getbbox()
                    pages[idx]["raster_width_px"] = page_image.width
                    pages[idx]["raster_height_px"] = page_image.height
                    pages[idx]["content_bbox"] = list(bbox) if bbox else None
                    if bbox is None:
                        issues.append({"code": "visually_blank_page", "page": idx + 1})
                    else:
                        left, top, right, bottom = bbox
                        edge_distance = min(left, top, page_image.width - right, page_image.height - bottom)
                        pages[idx]["min_content_edge_distance_px"] = edge_distance
                        if edge_distance <= 1:
                            pages[idx]["warning"] = "content_touches_page_edge"

    return {"status": "passed" if not issues else "failed", "page_count": page_count, "issues": issues, "pages": pages, "raster_status": raster_status, "visual_model_status": "not_configured"}
